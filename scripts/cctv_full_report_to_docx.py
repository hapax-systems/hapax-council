#!/usr/bin/env python3
"""Generate comprehensive DOCX with action plans + CCTV verdicts."""

from __future__ import annotations

import json
from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

PLANS_DIR = Path.home() / "Documents/Personal/20-projects/hapax-research"
COUNCIL_VERDICTS = PLANS_DIR / "datasets/epistemic-quality/cctv-council-synthesis-verdicts.jsonl"
ACTION_VERDICTS = PLANS_DIR / "cctv-action-plan-verdicts.jsonl"
OUTPUT = Path.home() / "gdrive-drop" / "cctv-full-action-report.docx"

PLAN_FILES = [
    ("Grounding Loop Repair", "cctv-action-plan-grounding-loop.md"),
    ("Governance Hardening", "cctv-action-plan-governance.md"),
    ("Research Program Repair", "cctv-action-plan-research.md"),
    ("Operator Trust + Narration", "cctv-action-plan-operator-trust.md"),
    ("Publication + Market Positioning", "cctv-action-plan-publication.md"),
    ("Routing + Config Contradictions", "cctv-action-plan-routing.md"),
]


def add_verdict_table(doc, verdicts):
    if not verdicts:
        doc.add_paragraph("No verdicts available.")
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Claim"
    hdr[1].text = "Status"
    hdr[2].text = "Evidence"
    hdr[3].text = "Counter-Ev"
    hdr[4].text = "Scope"
    for v in verdicts:
        row = table.add_row().cells
        row[0].text = v.get("claim_id", "?")
        row[1].text = v.get("convergence_status", "?").upper()
        scores = v.get("scores", {})
        row[2].text = str(scores.get("evidence_adequacy", "—"))
        row[3].text = str(scores.get("counter_evidence_resilience", "—"))
        row[4].text = str(scores.get("scope_honesty", "—"))


def build_report():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("CCTV Comprehensive Action Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}\n"
        "Source: 4 Model Council sessions + 6-model CCTV disconfirmation + 6 codebase research agents\n"
        "Protocol: 5-phase adversarial deliberation (blind scoring → evidence matrix → adversarial challenge → revision → convergence)\n"
        "Models: Claude Opus, Claude Sonnet, Gemini 3 Pro, Command-R 35B, Perplexity Sonar, Mistral Large"
    )

    # Part 1: Executive Summary
    doc.add_heading("Part 1: Executive Summary", level=1)

    council_verdicts = []
    if COUNCIL_VERDICTS.exists():
        council_verdicts = [json.loads(l) for l in COUNCIL_VERDICTS.read_text().splitlines() if l.strip()]

    action_verdicts = []
    if ACTION_VERDICTS.exists():
        action_verdicts = [json.loads(l) for l in ACTION_VERDICTS.read_text().splitlines() if l.strip()]

    doc.add_paragraph(
        f"Council synthesis claims evaluated: {len(council_verdicts)}\n"
        f"Action plan claims evaluated: {len(action_verdicts)}\n"
        f"Total CCTV verdicts: {len(council_verdicts) + len(action_verdicts)}"
    )

    # Part 2: CCTV Verdicts on Council Synthesis (Councils II-III)
    doc.add_heading("Part 2: CCTV Verdicts on Council Synthesis Claims", level=1)
    add_verdict_table(doc, council_verdicts)

    for v in council_verdicts:
        doc.add_heading(f"{v['claim_id']}: {v['convergence_status'].upper()}", level=3)
        p = doc.add_paragraph()
        p.add_run("Claim: ").bold = True
        p.add_run(v.get("claim_text", "")[:300])
        if v.get("research_findings"):
            p = doc.add_paragraph()
            p.add_run("Research: ").bold = True
            for f in v["research_findings"][:5]:
                doc.add_paragraph(f[:150], style="List Bullet")
        if v.get("adversarial_exchanges"):
            p = doc.add_paragraph()
            p.add_run("Adversarial: ").bold = True
            for e in v["adversarial_exchanges"][:3]:
                doc.add_paragraph(
                    f"{e['axis']}: {e.get('response_text', '')[:200]}",
                    style="List Bullet",
                )

    doc.add_page_break()

    # Part 3: CCTV Verdicts on Action Plans
    doc.add_heading("Part 3: CCTV Verdicts on Action Plan Claims", level=1)
    add_verdict_table(doc, action_verdicts)

    for v in action_verdicts:
        doc.add_heading(f"{v['claim_id']} ({v.get('domain','')}): {v['convergence_status'].upper()}", level=3)
        p = doc.add_paragraph()
        p.add_run("Claim: ").bold = True
        p.add_run(v.get("claim_text", "")[:300])
        if v.get("disagreement_log"):
            p = doc.add_paragraph()
            p.add_run("Disagreements: ").bold = True
            for d in v["disagreement_log"][:3]:
                doc.add_paragraph(d[:200], style="List Bullet")

    doc.add_page_break()

    # Part 4: Full Action Plans
    doc.add_heading("Part 4: Full Action Plans (6 domains)", level=1)

    for plan_title, plan_file in PLAN_FILES:
        plan_path = PLANS_DIR / plan_file
        if not plan_path.exists():
            continue

        doc.add_heading(plan_title, level=2)
        content = plan_path.read_text(encoding="utf-8")

        for line in content.split("\n"):
            line = line.rstrip()
            if line.startswith("# "):
                doc.add_heading(line[2:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=4)
            elif line.startswith("### "):
                p = doc.add_paragraph()
                p.add_run(line[4:]).bold = True
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("| "):
                doc.add_paragraph(line, style="No Spacing")
            elif line.strip():
                doc.add_paragraph(line)

        doc.add_page_break()

    # Part 5: Input Documents Reference
    doc.add_heading("Part 5: Source Documents", level=1)
    doc.add_paragraph(
        "Council Synthesis II: Trust, Perspectival Honesty, RLHF Trust Deficit\n"
        "Council Synthesis III: Non-RLHF Grounding Models, Stigmergic Architecture\n"
        "Council Synthesis IV: Repair and Leverage Roadmap\n"
        "CCTV Evidence Verification: 10/12 code citations confirmed\n"
        "CCTV Adversarial Pass: 6 claims attacked, all narrowed\n"
        "CCTV Actionable Findings: 21 items across 4 priority tiers"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Report saved: {OUTPUT} ({OUTPUT.stat().st_size / 1024:.0f} KB)")


if __name__ == "__main__":
    build_report()
